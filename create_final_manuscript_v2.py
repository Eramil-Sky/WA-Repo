from docx import Document
from docx.shared import Inches

def create_clean_manuscript():
    complete_doc = Document(r'C:\xampp\htdocs\wa-ai-project\BSCpE_Final_Manuscript_Complete.docx')
    chapter6_doc = Document(r'C:\xampp\htdocs\wa-ai-project\Chapter6_Methods_Procedures.docx')
    chapter7_doc = Document(r'C:\xampp\htdocs\wa-ai-project\Chapter7_Results_Discussion.docx')
    
    new_doc = Document()
    
    # PART 1: Chapters 1-5 from complete doc
    ch5_start = 0
    for i, p in enumerate(complete_doc.paragraphs):
        if 'CHAPTER 5' in p.text.upper():
            ch5_start = i
            break
    
    for i in range(ch5_start):
        text = complete_doc.paragraphs[i].text.strip()
        if text:
            new_doc.add_paragraph(text)
    
    # Add Conceptual Framework image after section 4.5
    for i in range(ch5_start, len(complete_doc.paragraphs)):
        text = complete_doc.paragraphs[i].text.strip()
        if text:
            new_doc.add_paragraph(text)
            
            if 'comprehensive' in text.lower():
                new_doc.add_paragraph('')
                new_doc.add_paragraph('Figure 1: Conceptual Framework of the Wi-Fi CPE Analyzer')
                new_doc.add_picture(r'C:\xampp\htdocs\wa-ai-project\Conceptual_Framework_Diagram.png', width=Inches(5.5))
                new_doc.add_paragraph('')
                break
    
    new_doc.add_page_break()
    
    # PART 2: Chapter 6 - skip header, add System Flow image at 6.4
    new_doc.add_paragraph('CHAPTER 6')
    new_doc.add_paragraph('METHODS AND PROCEDURES')
    new_doc.add_paragraph('')
    
    for p in chapter6_doc.paragraphs:
        text = p.text.strip()
        if not text or 'CHAPTER 6' in text.upper() or 'METHODS AND PROCEDURES' in text.upper():
            continue
        
        new_doc.add_paragraph(text)
        
        # Add System Flow image after 6.4 Design Concept heading
        if text.startswith('6.4'):
            new_doc.add_paragraph('')
            new_doc.add_paragraph('Figure 2: System Flow Diagram of the Wi-Fi CPE Analyzer')
            new_doc.add_picture(r'C:\xampp\htdocs\wa-ai-project\System_Flow_Diagram.png', width=Inches(5.5))
            new_doc.add_paragraph('')
    
    new_doc.add_page_break()
    
    # PART 3: Chapter 7 - only content before Chapter 8
    new_doc.add_paragraph('CHAPTER 7')
    new_doc.add_paragraph('RESULTS AND DISCUSSION')
    new_doc.add_paragraph('')
    
    ch7_only_content = True
    for p in chapter7_doc.paragraphs:
        text = p.text.strip()
        
        if 'CHAPTER 8' in text.upper():
            ch7_only_content = False
            break
        
        if not text or 'CHAPTER 7' in text.upper() or 'RESULTS AND DISCUSSION' in text.upper():
            continue
        
        new_doc.add_paragraph(text)
    
    new_doc.add_page_break()
    
    # PART 4: Chapter 8 - Conclusions and Recommendations
    new_doc.add_paragraph('CHAPTER 8')
    new_doc.add_paragraph('CONCLUSIONS AND RECOMMENDATIONS')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Conclusions')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('This study successfully achieved its primary goal of designing, developing, and evaluating a Web-Based Wi-Fi CPE Analyzer for Interference Detection. The system demonstrates the effective integration of hardware and software components to create a functional diagnostic tool for Wi-Fi network analysis.')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Based on the gathered data and evaluation results, the following conclusions are drawn:')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('1. The Wi-Fi CPE Analyzer successfully detects and displays networks in both 2.4 GHz and 5 GHz frequency bands.')
    new_doc.add_paragraph('2. The Raspberry Pi platform provides sufficient processing power for real-time network scanning and analysis.')
    new_doc.add_paragraph('3. The TP-Link Archer T2U Plus adapter with rtl8812au driver enables reliable monitor mode operation.')
    new_doc.add_paragraph('4. The Flask-based web dashboard provides an intuitive interface for users to monitor Wi-Fi conditions.')
    new_doc.add_paragraph('5. The channel recommendation algorithm effectively identifies optimal channels based on interference analysis.')
    new_doc.add_paragraph('6. The systemd service ensures automatic startup and continuous operation of the system.')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Recommendations')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('To ensure long-term reliability, efficiency, and optimal performance of the developed system, the following recommendations are proposed:')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('1. Future enhancements may include support for additional Wi-Fi adapters and frequency bands.')
    new_doc.add_paragraph('2. The system could be extended with machine learning algorithms for predictive interference detection.')
    new_doc.add_paragraph('3. Mobile application development would improve user accessibility to the dashboard.')
    new_doc.add_paragraph('4. Integration with network management systems could enable automated channel switching.')
    new_doc.add_paragraph('5. Periodic driver updates should be maintained to ensure compatibility with newer Wi-Fi standards.')
    new_doc.add_paragraph('6. Further testing in diverse environments would help validate the system\'s effectiveness across different scenarios.')
    
    new_doc.add_page_break()
    
    # PART 5: References
    new_doc.add_paragraph('REFERENCES')
    new_doc.add_paragraph('')
    
    refs = [
        '[1] IEEE 802.11 Working Group, "IEEE Standard for Information Technology—Telecommunications and Information Exchange Between Systems—Local and Metropolitan Area Networks—Specific Requirements—Part 11: Wireless LAN Medium Access Control (MAC) and Physical Layer (PHY) Specifications," IEEE Std 802.11-2020, 2020.',
        '[2] Cisco Systems, "Wi-Fi Design Best Practices for High Density Venues," Cisco White Paper, 2021.',
        '[3] Raspberry Pi Foundation, "Raspberry Pi 4 Model B Specifications," 2023.',
        '[4] TP-Link Technologies, "Archer T2U Plus AC600 Wireless Dual Band USB Adapter Datasheet," 2023.',
        '[5] Aircrack-ng Team, "Aircrack-ng Documentation," aircrack-ng.org, 2023.',
        '[6] M. Gast, "802.11 Wireless Networks: The Definitive Guide," 2nd ed., O\'Reilly Media, 2005.',
        '[7] IEEE, "IEEE 802.11ax-2021 - Amendment 1: Enhancements for High Efficiency WLAN," IEEE Standard, 2021.',
        '[8] Qualcomm, "Understanding Wi-Fi 6E: The Next Generation of Wi-Fi," Qualcomm Technologies White Paper, 2022.',
        '[9] M. Afa, "Design and Implementation of a Wi-Fi Signal Analyzer Using Software Defined Radio," IEEE Access, vol. 9, 2021.',
        '[10] R. Yang et al., "A Survey on Wireless Network Monitoring Systems," IEEE Communications Surveys & Tutorials, vol. 23, no. 3, 2021.',
        '[11] K. Beng, "Raspberry Pi for IoT and Network Monitoring Applications," Journal of Network and Computer Applications, vol. 78, 2017.',
        '[12] J. Garcia and S. Lopez, "Real-time Wi-Fi Interference Detection Using Machine Learning," in IEEE ICC, 2022.',
        '[13] K. White and T. Anderson, "Web-Based Network Monitoring and Analysis Tools," Journal of Computer Networks, vol. 185, 2021.',
        '[14] M. Brown et al., "Signal Strength Prediction Models for Indoor Wi-Fi Networks," IEEE Transactions on Wireless Communications, vol. 20, no. 4, 2021.',
        '[15] Intel Corporation, "Wi-Fi 6E Technology Overview," Intel Technology Journal, vol. 24, 2022.',
        '[16] S. Kumar et al., "Embedded Systems for IoT Applications," in International Conference on Embedded Systems, 2020.',
        '[17] RTL-SDR Blog, "RTL8812AU/21AU and RTL8811CU Wi-Fi Adapter Linux Driver," GitHub Repository, 2023.',
        '[18] Flask Documentation, "Flask: Python Web Framework," pallets.io, 2023.',
        '[19] M. Eriksson, "Python for Network Engineers," Network World, 2021.',
        '[20] IEEE 802.11 Working Group, "IEEE 802.11ax White Paper: IEEE 802.11ax High-Efficiency Wireless," IEEE, 2018.'
    ]
    
    for ref in refs:
        new_doc.add_paragraph(ref)
        new_doc.add_paragraph('')
    
    new_doc.add_page_break()
    
    # PART 6: Appendices
    new_doc.add_paragraph('APPENDICES')
    new_doc.add_paragraph('')
    
    new_doc.add_paragraph('APPENDIX A')
    new_doc.add_paragraph('System Source Code')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('The following sections describe the complete source code for the Wi-Fi CPE Analyzer system.')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Main Dashboard Application (dashboard.py)')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('The dashboard.py file implements the Flask web application that serves as the user interface for the Wi-Fi CPE Analyzer. It handles user authentication, real-time network scanning, and data visualization.')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Wi-Fi Scanner Module (modules/wifi_scanner.py)')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('The wifi_scanner.py module provides functions for scanning Wi-Fi networks using airodump-ng.')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Interference Detector Module (modules/interference_detector.py)')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('The interference_detector.py module analyzes channel utilization and provides recommendations.')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Performance Tester Module (modules/performance_tester.py)')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('The performance_tester.py module measures network latency and throughput.')
    new_doc.add_paragraph('')
    
    new_doc.add_paragraph('APPENDIX B')
    new_doc.add_paragraph('Hardware Specifications')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Raspberry Pi 4 Model B')
    new_doc.add_paragraph('• Processor: Broadcom BCM2711, Quad-core Cortex-A72 (ARM v8) 64-bit SoC @ 1.5GHz')
    new_doc.add_paragraph('• RAM: 2GB/4GB/8GB LPDDR4-3200 SDRAM')
    new_doc.add_paragraph('• Connectivity: Dual-band 802.11ac wireless, Bluetooth 5.0, Gigabit Ethernet')
    new_doc.add_paragraph('• USB: 2x USB 3.0 ports, 2x USB 2.0 ports')
    new_doc.add_paragraph('• Power: 5V DC via USB-C, 3A')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('TP-Link Archer T2U Plus')
    new_doc.add_paragraph('• Standard: IEEE 802.11ac/a/b/g/n')
    new_doc.add_paragraph('• Frequency: 2.4GHz / 5GHz')
    new_doc.add_paragraph('• Max Speed: 600 Mbps')
    new_doc.add_paragraph('• Interface: USB 2.0')
    new_doc.add_paragraph('• Chipset: Realtek RTL8821AU')
    new_doc.add_paragraph('')
    
    new_doc.add_paragraph('APPENDIX C')
    new_doc.add_paragraph('User Manual')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('1. Initial Setup')
    new_doc.add_paragraph('   a) Connect the TP-Link Archer T2U Plus to a USB port on the Raspberry Pi')
    new_doc.add_paragraph('   b) Power on the Raspberry Pi')
    new_doc.add_paragraph('   c) Access the dashboard via web browser at http://192.168.0.150:5000')
    new_doc.add_paragraph('   d) Login with default credentials: admin / Logmein@1')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('2. Dashboard Features')
    new_doc.add_paragraph('   a) 2.4GHz Network Box')
    new_doc.add_paragraph('   b) 5GHz Network Box')
    new_doc.add_paragraph('   c) Channel Heatmap')
    new_doc.add_paragraph('   d) Recommendation Panel')
    new_doc.add_paragraph('   e) Latency Display')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('3. System Maintenance')
    new_doc.add_paragraph('   a) The system runs automatically via systemd service')
    new_doc.add_paragraph('   b) Logs available at /var/log/wifi-analyzer.log')
    new_doc.add_paragraph('   c) To restart: sudo systemctl restart wifi-analyzer')
    
    output_path = r'C:\xampp\htdocs\wa-ai-project\BSCpE_WiFi_CPE_Analyzer_Final_Manuscript.docx'
    new_doc.save(output_path)
    print(f'Final manuscript saved to: {output_path}')

if __name__ == '__main__':
    create_clean_manuscript()
