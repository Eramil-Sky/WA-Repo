from docx import Document
from docx.shared import Inches

def add_conceptual_framework_image():
    doc = Document(r'C:\xampp\htdocs\wa-ai-project\BSCpE_Final_Manuscript_v9_Final.docx')
    
    new_doc = Document()
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        new_doc.add_paragraph(text)
        
        # Add Conceptual Framework image after Conceptual Framework section
        if 'OUTPUT STAGE' in text.upper():
            new_doc.add_paragraph('')
            new_doc.add_paragraph('Figure 1: Conceptual Framework of the Wi-Fi CPE Analyzer')
            new_doc.add_picture(r'C:\xampp\htdocs\wa-ai-project\Conceptual_Framework_Diagram.png', width=Inches(5.5))
            new_doc.add_paragraph('')
    
    output_path = r'C:\xampp\htdocs\wa-ai-project\BSCpE_Final_Manuscript_Complete_v10.docx'
    new_doc.save(output_path)
    print(f"Manuscript with Conceptual Framework image saved to: {output_path}")

if __name__ == '__main__':
    add_conceptual_framework_image()
